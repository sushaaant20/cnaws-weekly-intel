from io import BytesIO
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    Document = None
    Inches = None


def _format_number(value, decimals=0):
    if value is None:
        return "-"
    if decimals == 0:
        return f"{int(round(value)):,}"
    return f"{value:,.{decimals}f}"


def _format_pct(value):
    if value is None:
        return "New"
    return f"{value:+.1%}"


def _build_table_data(dataframe, columns):
    header = [label for _, label in columns]
    rows = [header]

    if dataframe.empty:
        rows.append(["No data"] + [""] * (len(columns) - 1))
        return rows

    for _, row in dataframe.iterrows():
        rows.append([row[column] for column, _ in columns])

    return rows


def _executive_summary_lines(report_context):
    return [line for line in report_context["executive_summary"].splitlines() if line.strip()]


def _summary_lines(report_context):
    score = report_context["intelligence_score"]
    actor = report_context["actor_metrics"]
    return [
        f"Intelligence Score: {_format_number(score['final_score'])} / 100",
        f"Risk Level: {score['risk_level']}",
        f"Security Force Ratio: {actor['ct_ratio']:.2f} ({actor['interpretation']})",
    ]


def _score_driver_lines(report_context):
    return report_context["intelligence_score"]["drivers"]


def build_pdf_report(report_context):
    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontSize=18,
        textColor=colors.HexColor("#102a43"),
        spaceAfter=6,
    )
    section_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=colors.HexColor("#0f172a"),
        spaceAfter=6,
    )
    body_style = styles["BodyText"]

    story = [
        Paragraph("CNAWS Weekly Intelligence Dashboard", title_style),
        Paragraph(report_context["window_header"], body_style),
        Spacer(1, 12),
    ]

    story.append(Paragraph("Executive Summary", section_style))
    for line in _executive_summary_lines(report_context):
        story.append(Paragraph(line, body_style))
    story.append(Spacer(1, 12))

    story.append(Paragraph("Intelligence Score", section_style))
    for line in _summary_lines(report_context):
        story.append(Paragraph(line, body_style))
    for driver in _score_driver_lines(report_context):
        story.append(Paragraph(f"- {driver}", body_style))
    story.append(Spacer(1, 12))

    metric_rows = [["Metric", "Current", "Previous", "Delta"]]
    for label, metric in report_context["metric_cards"]:
        metric_rows.append(
            [
                label,
                _format_number(metric["current"], 1 if isinstance(metric["current"], float) and not float(metric["current"]).is_integer() else 0),
                _format_number(metric["previous"], 1 if isinstance(metric["previous"], float) and not float(metric["previous"]).is_integer() else 0),
                _format_number(metric["delta"], 1 if isinstance(metric["delta"], float) and not float(metric["delta"]).is_integer() else 0),
            ]
        )

    metric_table = Table(metric_rows, repeatRows=1)
    metric_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#102a43")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d9e2ec")),
                ("BACKGROUND", (0, 1), (-1, -1), colors.HexColor("#f8fafc")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([Paragraph("Key Metrics", section_style), metric_table, Spacer(1, 12)])

    actor_rows = [
        ["Metric", "Current", "Previous", "Delta"],
        [
            "Militant Incidents",
            report_context["actor_metrics"]["militant_incidents"]["current"],
            report_context["actor_metrics"]["militant_incidents"]["previous"],
            report_context["actor_metrics"]["militant_incidents"]["delta"],
        ],
        [
            "Security Force Operations",
            report_context["actor_metrics"]["ct_operations"]["current"],
            report_context["actor_metrics"]["ct_operations"]["previous"],
            report_context["actor_metrics"]["ct_operations"]["delta"],
        ],
        ["CT Ratio", f"{report_context['actor_metrics']['ct_ratio']:.2f}", "-", report_context["actor_metrics"]["interpretation"]],
    ]
    actor_table = Table(actor_rows, repeatRows=1)
    actor_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e2ec")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bcccdc")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("PADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([Paragraph("Actor Separation", section_style), actor_table, Spacer(1, 12)])

    if report_context["map_image"] is not None:
        story.append(Paragraph("Geographic View", section_style))
        story.append(RLImage(BytesIO(report_context["map_image"]), width=6.6 * inch, height=3.9 * inch))
        story.append(Spacer(1, 12))

    district_rows = _build_table_data(
        report_context["district_breakdown"].head(10).assign(
            pct_change=lambda frame: frame["pct_change"].apply(_format_pct),
            share=lambda frame: (frame["share"] * 100).round(1).astype(str) + "%",
            lethality=lambda frame: frame["lethality"].round(2),
        ),
        [
            ("district", "District"),
            ("current", "Current"),
            ("previous", "Previous"),
            ("delta", "Delta"),
            ("pct_change", "% Change"),
            ("share", "Share"),
            ("lethality", "Lethality"),
        ],
    )

    event_rows = _build_table_data(
        report_context["event_breakdown"].head(10).assign(
            pct_change=lambda frame: frame["pct_change"].apply(_format_pct),
            share=lambda frame: (frame["share"] * 100).round(1).astype(str) + "%",
        ),
        [
            ("event_type", "Event Type"),
            ("current", "Current"),
            ("previous", "Previous"),
            ("delta", "Delta"),
            ("pct_change", "% Change"),
            ("share", "Share"),
        ],
    )

    for heading, rows in (
        ("District Intelligence", district_rows),
        ("Event Analysis", event_rows),
    ):
        table = Table(rows, repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e2ec")),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bcccdc")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("PADDING", (0, 0), (-1, -1), 5),
                ]
            )
        )
        story.extend([Paragraph(heading, section_style), table, Spacer(1, 12)])

    expansion = report_context["expansion_analysis"]
    story.append(Paragraph("Expansion Panel", section_style))
    story.append(
        Paragraph(
            f"New Districts: {expansion['new_districts']} | Expansion Index: {expansion['expansion_index']:.2f} | {expansion['tag']}",
            body_style,
        )
    )
    story.append(Spacer(1, 6))
    expansion_rows = _build_table_data(
        expansion["table"],
        [("district", "District"), ("incidents", "Incidents"), ("casualties", "Casualties")],
    )
    expansion_table = Table(expansion_rows, repeatRows=1)
    expansion_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e2ec")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bcccdc")),
                ("PADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([expansion_table, Spacer(1, 12)])

    story.append(Paragraph("Tactical Insights", section_style))
    tactical = report_context["tactical_shift"]
    if tactical["top_increase"] is not None:
        story.append(
            Paragraph(
                f"Top Increase: {tactical['top_increase']['event_type']} ({_format_pct(tactical['top_increase']['pct_change'])})",
                body_style,
            )
        )
    if tactical["top_decrease"] is not None:
        story.append(
            Paragraph(
                f"Top Decrease: {tactical['top_decrease']['event_type']} ({_format_pct(tactical['top_decrease']['pct_change'])})",
                body_style,
            )
        )
    for observation in tactical["observations"]:
        story.append(Paragraph(f"- {observation}", body_style))
    story.append(Paragraph(f"Interpretation: {tactical['interpretation']}", body_style))
    story.append(Spacer(1, 12))

    high_impact = report_context["high_impact_incidents"].copy()
    if not high_impact.empty:
        high_impact["date"] = pd.to_datetime(high_impact["date"]).dt.strftime("%d %b %Y")
        high_impact = high_impact.rename(columns={"casualties_total": "casualties"})
    high_impact_rows = _build_table_data(
        high_impact,
        [("date", "Date"), ("district", "District"), ("event_type", "Event Type"), ("casualties", "Casualties")],
    )
    high_impact_table = Table(high_impact_rows, repeatRows=1)
    high_impact_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e2ec")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bcccdc")),
                ("PADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.extend([Paragraph("High Impact Incidents", section_style), high_impact_table])

    document.build(story)
    return buffer.getvalue()


def build_docx_report(report_context):
    if Document is not None:
        return _build_docx_with_python_docx(report_context)
    return _build_docx_fallback(report_context)


def _build_docx_with_python_docx(report_context):
    document = Document()
    document.add_heading("CNAWS Weekly Intelligence Dashboard", level=0)
    document.add_paragraph(report_context["window_header"])

    document.add_heading("Executive Summary", level=1)
    for line in _executive_summary_lines(report_context):
        document.add_paragraph(line)

    document.add_heading("Intelligence Score", level=1)
    for line in _summary_lines(report_context):
        document.add_paragraph(line)
    for driver in _score_driver_lines(report_context):
        document.add_paragraph(driver, style="List Bullet")

    document.add_heading("Key Metrics", level=1)
    _add_metric_table(document, report_context["metric_cards"])

    document.add_heading("Actor Separation", level=1)
    actor_df = pd.DataFrame(
        [
            {
                "Metric": "Militant Incidents",
                "Current": report_context["actor_metrics"]["militant_incidents"]["current"],
                "Previous": report_context["actor_metrics"]["militant_incidents"]["previous"],
                "Delta": report_context["actor_metrics"]["militant_incidents"]["delta"],
            },
            {
                "Metric": "Security Force Operations",
                "Current": report_context["actor_metrics"]["ct_operations"]["current"],
                "Previous": report_context["actor_metrics"]["ct_operations"]["previous"],
                "Delta": report_context["actor_metrics"]["ct_operations"]["delta"],
            },
            {
                "Metric": "CT Ratio",
                "Current": f"{report_context['actor_metrics']['ct_ratio']:.2f}",
                "Previous": "-",
                "Delta": report_context["actor_metrics"]["interpretation"],
            },
        ]
    )
    _add_dataframe_section(
        document,
        None,
        actor_df,
        [("Metric", "Metric"), ("Current", "Current"), ("Previous", "Previous"), ("Delta", "Delta")],
    )

    if report_context["map_image"] is not None and Inches is not None:
        document.add_heading("Geographic View", level=1)
        document.add_picture(BytesIO(report_context["map_image"]), width=Inches(6.3))

    _add_dataframe_section(
        document,
        "District Intelligence",
        report_context["district_breakdown"].head(10).assign(
            pct_change=lambda frame: frame["pct_change"].apply(_format_pct),
            share=lambda frame: (frame["share"] * 100).round(1).astype(str) + "%",
        ),
        [
            ("district", "District"),
            ("current", "Current"),
            ("previous", "Previous"),
            ("delta", "Delta"),
            ("pct_change", "% Change"),
            ("share", "Share"),
            ("lethality", "Lethality"),
        ],
    )
    _add_dataframe_section(
        document,
        "Event Analysis",
        report_context["event_breakdown"].head(10).assign(
            pct_change=lambda frame: frame["pct_change"].apply(_format_pct),
            share=lambda frame: (frame["share"] * 100).round(1).astype(str) + "%",
        ),
        [
            ("event_type", "Event Type"),
            ("current", "Current"),
            ("previous", "Previous"),
            ("delta", "Delta"),
            ("pct_change", "% Change"),
            ("share", "Share"),
        ],
    )

    document.add_heading("Expansion Panel", level=1)
    expansion = report_context["expansion_analysis"]
    document.add_paragraph(
        f"New Districts: {expansion['new_districts']} | Expansion Index: {expansion['expansion_index']:.2f} | {expansion['tag']}"
    )
    _add_dataframe_section(
        document,
        None,
        expansion["table"],
        [("district", "District"), ("incidents", "Incidents"), ("casualties", "Casualties")],
    )

    document.add_heading("Tactical Insights", level=1)
    tactical = report_context["tactical_shift"]
    if tactical["top_increase"] is not None:
        document.add_paragraph(
            f"Top Increase: {tactical['top_increase']['event_type']} ({_format_pct(tactical['top_increase']['pct_change'])})"
        )
    if tactical["top_decrease"] is not None:
        document.add_paragraph(
            f"Top Decrease: {tactical['top_decrease']['event_type']} ({_format_pct(tactical['top_decrease']['pct_change'])})"
        )
    for observation in tactical["observations"]:
        document.add_paragraph(observation, style="List Bullet")
    document.add_paragraph(f"Interpretation: {tactical['interpretation']}")

    high_impact = report_context["high_impact_incidents"].copy()
    if not high_impact.empty:
        high_impact["date"] = pd.to_datetime(high_impact["date"]).dt.strftime("%d %b %Y")
        high_impact = high_impact.rename(columns={"casualties_total": "casualties"})
    _add_dataframe_section(
        document,
        "High Impact Incidents",
        high_impact,
        [("date", "Date"), ("district", "District"), ("event_type", "Event Type"), ("casualties", "Casualties")],
    )

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _add_metric_table(document, metric_cards):
    table = document.add_table(rows=1, cols=4)
    headers = ("Metric", "Current", "Previous", "Delta")
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    for label, metric in metric_cards:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = _format_number(metric["current"], 1 if isinstance(metric["current"], float) and not float(metric["current"]).is_integer() else 0)
        row[2].text = _format_number(metric["previous"], 1 if isinstance(metric["previous"], float) and not float(metric["previous"]).is_integer() else 0)
        row[3].text = _format_number(metric["delta"], 1 if isinstance(metric["delta"], float) and not float(metric["delta"]).is_integer() else 0)


def _add_dataframe_section(document, heading, dataframe, columns):
    if heading:
        document.add_heading(heading, level=1)

    table = document.add_table(rows=1, cols=len(columns))
    for index, (_, label) in enumerate(columns):
        table.rows[0].cells[index].text = label

    if dataframe.empty:
        row = table.add_row().cells
        row[0].text = "No data"
        return

    for _, series in dataframe.iterrows():
        row = table.add_row().cells
        for index, (column, _) in enumerate(columns):
            row[index].text = str(series[column])


def _build_docx_fallback(report_context):
    body_parts = []
    media_entries = []

    def add_paragraph(text, bold=False):
        if text is None:
            return
        if bold:
            run = f"<w:r><w:rPr><w:b/></w:rPr><w:t xml:space='preserve'>{escape(str(text))}</w:t></w:r>"
        else:
            run = f"<w:r><w:t xml:space='preserve'>{escape(str(text))}</w:t></w:r>"
        body_parts.append(f"<w:p>{run}</w:p>")

    def add_table(dataframe, columns):
        rows = _build_table_data(dataframe, columns)
        column_count = len(columns)
        grid = "".join("<w:gridCol w:w='2200'/>" for _ in range(column_count))
        table_rows = []
        for row_index, row in enumerate(rows):
            cells = []
            for value in row:
                run_props = "<w:rPr><w:b/></w:rPr>" if row_index == 0 else ""
                cells.append(
                    "<w:tc>"
                    "<w:tcPr><w:tcW w:w='2200' w:type='dxa'/></w:tcPr>"
                    f"<w:p><w:r>{run_props}<w:t xml:space='preserve'>{escape(str(value))}</w:t></w:r></w:p>"
                    "</w:tc>"
                )
            table_rows.append(f"<w:tr>{''.join(cells)}</w:tr>")

        body_parts.append(
            "<w:tbl>"
            "<w:tblPr><w:tblW w:w='0' w:type='auto'/>"
            "<w:tblBorders>"
            "<w:top w:val='single' w:sz='4' w:space='0' w:color='D9E2EC'/>"
            "<w:left w:val='single' w:sz='4' w:space='0' w:color='D9E2EC'/>"
            "<w:bottom w:val='single' w:sz='4' w:space='0' w:color='D9E2EC'/>"
            "<w:right w:val='single' w:sz='4' w:space='0' w:color='D9E2EC'/>"
            "<w:insideH w:val='single' w:sz='4' w:space='0' w:color='D9E2EC'/>"
            "<w:insideV w:val='single' w:sz='4' w:space='0' w:color='D9E2EC'/>"
            "</w:tblBorders></w:tblPr>"
            f"<w:tblGrid>{grid}</w:tblGrid>"
            f"{''.join(table_rows)}"
            "</w:tbl>"
        )

    def add_image(image_bytes, image_name):
        if image_bytes is None:
            return

        image = PILImage.open(BytesIO(image_bytes))
        width_px, height_px = image.size
        max_width_px = 610
        if width_px > max_width_px:
            scale = max_width_px / width_px
            width_px = int(width_px * scale)
            height_px = int(height_px * scale)

        cx = width_px * 9525
        cy = height_px * 9525
        rel_id = f"rIdImage{len(media_entries) + 1}"
        media_name = f"media/{image_name}"
        media_entries.append((rel_id, media_name, image_bytes))

        body_parts.append(
            "<w:p><w:r><w:drawing>"
            "<wp:inline distT='0' distB='0' distL='0' distR='0'>"
            f"<wp:extent cx='{cx}' cy='{cy}'/>"
            "<wp:effectExtent l='0' t='0' r='0' b='0'/>"
            f"<wp:docPr id='{len(media_entries)}' name='{escape(image_name)}'/>"
            "<wp:cNvGraphicFramePr><a:graphicFrameLocks noChangeAspect='1'/></wp:cNvGraphicFramePr>"
            "<a:graphic>"
            "<a:graphicData uri='http://schemas.openxmlformats.org/drawingml/2006/picture'>"
            "<pic:pic>"
            "<pic:nvPicPr>"
            f"<pic:cNvPr id='{len(media_entries)}' name='{escape(image_name)}'/>"
            "<pic:cNvPicPr/>"
            "</pic:nvPicPr>"
            "<pic:blipFill>"
            f"<a:blip r:embed='{rel_id}'/>"
            "<a:stretch><a:fillRect/></a:stretch>"
            "</pic:blipFill>"
            "<pic:spPr>"
            "<a:xfrm><a:off x='0' y='0'/>"
            f"<a:ext cx='{cx}' cy='{cy}'/></a:xfrm>"
            "<a:prstGeom prst='rect'><a:avLst/></a:prstGeom>"
            "</pic:spPr>"
            "</pic:pic>"
            "</a:graphicData>"
            "</a:graphic>"
            "</wp:inline>"
            "</w:drawing></w:r></w:p>"
        )

    add_paragraph("CNAWS Weekly Intelligence Dashboard", bold=True)
    add_paragraph(report_context["window_header"])

    add_paragraph("Executive Summary", bold=True)
    for line in _executive_summary_lines(report_context):
        add_paragraph(line)

    add_paragraph("Intelligence Score", bold=True)
    for line in _summary_lines(report_context):
        add_paragraph(line)
    for driver in _score_driver_lines(report_context):
        add_paragraph(f"- {driver}")

    add_paragraph("Key Metrics", bold=True)
    metric_df = pd.DataFrame(
        [
            {
                "Metric": label,
                "Current": metric["current"],
                "Previous": metric["previous"],
                "Delta": metric["delta"],
            }
            for label, metric in report_context["metric_cards"]
        ]
    )
    add_table(metric_df, [("Metric", "Metric"), ("Current", "Current"), ("Previous", "Previous"), ("Delta", "Delta")])

    add_paragraph("Actor Separation", bold=True)
    actor_df = pd.DataFrame(
        [
            {
                "Metric": "Militant Incidents",
                "Current": report_context["actor_metrics"]["militant_incidents"]["current"],
                "Previous": report_context["actor_metrics"]["militant_incidents"]["previous"],
                "Delta": report_context["actor_metrics"]["militant_incidents"]["delta"],
            },
            {
                "Metric": "Security Force Operations",
                "Current": report_context["actor_metrics"]["ct_operations"]["current"],
                "Previous": report_context["actor_metrics"]["ct_operations"]["previous"],
                "Delta": report_context["actor_metrics"]["ct_operations"]["delta"],
            },
            {
                "Metric": "CT Ratio",
                "Current": f"{report_context['actor_metrics']['ct_ratio']:.2f}",
                "Previous": "-",
                "Delta": report_context["actor_metrics"]["interpretation"],
            },
        ]
    )
    add_table(actor_df, [("Metric", "Metric"), ("Current", "Current"), ("Previous", "Previous"), ("Delta", "Delta")])

    add_paragraph("Geographic View", bold=True)
    if report_context["map_image"] is not None:
        add_image(report_context["map_image"], "map.png")
    else:
        add_paragraph("Map image unavailable.")

    add_paragraph("District Intelligence", bold=True)
    add_table(
        report_context["district_breakdown"].head(10).assign(
            pct_change=lambda frame: frame["pct_change"].apply(_format_pct),
            share=lambda frame: (frame["share"] * 100).round(1).astype(str) + "%",
        ),
        [
            ("district", "District"),
            ("current", "Current"),
            ("previous", "Previous"),
            ("delta", "Delta"),
            ("pct_change", "% Change"),
            ("share", "Share"),
            ("lethality", "Lethality"),
        ],
    )

    add_paragraph("Event Analysis", bold=True)
    add_table(
        report_context["event_breakdown"].head(10).assign(
            pct_change=lambda frame: frame["pct_change"].apply(_format_pct),
            share=lambda frame: (frame["share"] * 100).round(1).astype(str) + "%",
        ),
        [
            ("event_type", "Event Type"),
            ("current", "Current"),
            ("previous", "Previous"),
            ("delta", "Delta"),
            ("pct_change", "% Change"),
            ("share", "Share"),
        ],
    )

    add_paragraph("Expansion Panel", bold=True)
    expansion = report_context["expansion_analysis"]
    add_paragraph(
        f"New Districts: {expansion['new_districts']} | Expansion Index: {expansion['expansion_index']:.2f} | {expansion['tag']}"
    )
    add_table(
        expansion["table"],
        [("district", "District"), ("incidents", "Incidents"), ("casualties", "Casualties")],
    )

    add_paragraph("Tactical Insights", bold=True)
    tactical = report_context["tactical_shift"]
    if tactical["top_increase"] is not None:
        add_paragraph(
            f"Top Increase: {tactical['top_increase']['event_type']} ({_format_pct(tactical['top_increase']['pct_change'])})"
        )
    if tactical["top_decrease"] is not None:
        add_paragraph(
            f"Top Decrease: {tactical['top_decrease']['event_type']} ({_format_pct(tactical['top_decrease']['pct_change'])})"
        )
    for observation in tactical["observations"]:
        add_paragraph(f"- {observation}")
    add_paragraph(f"Interpretation: {tactical['interpretation']}")

    add_paragraph("High Impact Incidents", bold=True)
    high_impact = report_context["high_impact_incidents"].copy()
    if not high_impact.empty:
        high_impact["date"] = pd.to_datetime(high_impact["date"]).dt.strftime("%d %b %Y")
        high_impact = high_impact.rename(columns={"casualties_total": "casualties"})
    add_table(
        high_impact,
        [("date", "Date"), ("district", "District"), ("event_type", "Event Type"), ("casualties", "Casualties")],
    )

    body_parts.append("<w:sectPr><w:pgSz w:w='12240' w:h='15840'/><w:pgMar w:top='720' w:right='720' w:bottom='720' w:left='720'/></w:sectPr>")

    document_xml = (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        "<w:document "
        "xmlns:w='http://schemas.openxmlformats.org/wordprocessingml/2006/main' "
        "xmlns:r='http://schemas.openxmlformats.org/officeDocument/2006/relationships' "
        "xmlns:wp='http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing' "
        "xmlns:a='http://schemas.openxmlformats.org/drawingml/2006/main' "
        "xmlns:pic='http://schemas.openxmlformats.org/drawingml/2006/picture'>"
        f"<w:body>{''.join(body_parts)}</w:body></w:document>"
    )

    content_types_xml = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<Types xmlns='http://schemas.openxmlformats.org/package/2006/content-types'>
  <Default Extension='rels' ContentType='application/vnd.openxmlformats-package.relationships+xml'/>
  <Default Extension='xml' ContentType='application/xml'/>
  <Default Extension='png' ContentType='image/png'/>
  <Override PartName='/word/document.xml' ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml'/>
  <Override PartName='/docProps/core.xml' ContentType='application/vnd.openxmlformats-package.core-properties+xml'/>
  <Override PartName='/docProps/app.xml' ContentType='application/vnd.openxmlformats-officedocument.extended-properties+xml'/>
</Types>"""

    rels_xml = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>
  <Relationship Id='rId1' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument' Target='word/document.xml'/>
  <Relationship Id='rId2' Type='http://schemas.openxmlformats.org/package/2006/relationships/metadata/core-properties' Target='docProps/core.xml'/>
  <Relationship Id='rId3' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/extended-properties' Target='docProps/app.xml'/>
</Relationships>"""

    document_rels = ["<?xml version='1.0' encoding='UTF-8' standalone='yes'?>", "<Relationships xmlns='http://schemas.openxmlformats.org/package/2006/relationships'>"]
    for rel_id, media_name, _ in media_entries:
        document_rels.append(
            f"<Relationship Id='{rel_id}' Type='http://schemas.openxmlformats.org/officeDocument/2006/relationships/image' Target='{media_name}'/>"
        )
    document_rels.append("</Relationships>")

    core_xml = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<cp:coreProperties xmlns:cp='http://schemas.openxmlformats.org/package/2006/metadata/core-properties'
 xmlns:dc='http://purl.org/dc/elements/1.1/'
 xmlns:dcterms='http://purl.org/dc/terms/'
 xmlns:dcmitype='http://purl.org/dc/dcmitype/'
 xmlns:xsi='http://www.w3.org/2001/XMLSchema-instance'>
  <dc:title>CNAWS Weekly Intelligence Dashboard</dc:title>
  <dc:creator>Codex</dc:creator>
</cp:coreProperties>"""

    app_xml = """<?xml version='1.0' encoding='UTF-8' standalone='yes'?>
<Properties xmlns='http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
 xmlns:vt='http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'>
  <Application>Codex</Application>
</Properties>"""

    output = BytesIO()
    with ZipFile(output, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", content_types_xml)
        archive.writestr("_rels/.rels", rels_xml)
        archive.writestr("docProps/core.xml", core_xml)
        archive.writestr("docProps/app.xml", app_xml)
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/_rels/document.xml.rels", "".join(document_rels))
        for _, media_name, media_bytes in media_entries:
            archive.writestr(f"word/{media_name}", media_bytes)

    return output.getvalue()
